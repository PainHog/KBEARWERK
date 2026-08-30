"""Unit tests for the KBEARWERK service layer (no GUI needed).

Run:  python -m pytest -q tests      (or)      python tests/test_services.py
"""

import os
import tempfile

# Isolate all local storage (config, jobdata, billing, activity, outbox) in temp.
os.environ.setdefault("KBEARWERK_CONFIG_DIR", tempfile.mkdtemp(prefix="kbw_cfg_"))

from kbearwerk import config
from kbearwerk.services import (files, excel, reliable, outbox, search, checklist,
                                billing, jobdata, activity, fieldscan, extract,
                                templates)


def _work():
    return tempfile.mkdtemp(prefix="kbw_work_")


def test_config_roundtrip():
    cfg = config.default_config()
    assert cfg["project_folder_pattern"] == "{number} - {name}"
    cfg["base_folder"] = "/tmp/x"
    config.save(cfg)
    assert config.load()["base_folder"] == "/tmp/x"


def test_folder_create_idempotent():
    work = _work()
    name = files.format_folder_name("{number} - {name}", number="2025-1", name="Smith: House?")
    assert ":" not in name and "?" not in name  # cleaned
    r1 = files.create_project_folder(work, name, ["01 - A", "02 - B"])
    assert r1.created and len(r1.subfolders_created) == 2
    r2 = files.create_project_folder(work, name, ["01 - A", "02 - B"])
    assert not r2.created and r2.subfolders_created == []  # nothing re-created


def test_excel_append_next_empty_and_update():
    work = _work()
    path = os.path.join(work, "list.xlsx")
    excel.ensure_workbook(path, ["Project Number", "Client", "Status"])
    excel.append_row(path, {"Project Number": "1", "Client": "A", "Status": "Open"})
    excel.append_row(path, {"Project Number": "2", "Client": "B", "Status": "Open"})
    # Same key updates, doesn't add a third row.
    res = excel.append_row(path, {"Project Number": "1", "Status": "Accepted"}, unique_key="Project Number")
    assert res == "updated"
    rows = excel.read_rows(path)
    assert len(rows) == 2
    row1 = next(r for r in rows if str(r["Project Number"]) == "1")
    assert row1["Status"] == "Accepted" and row1["Client"] == "A"  # other cells preserved


def test_excel_unavailable_is_classified():
    missing = os.path.join(_work(), "nope", "list.xlsx")  # parent dir absent
    try:
        excel.append_row(missing, {"x": 1})
        assert False, "should have raised"
    except excel.ExcelUnavailableError:
        pass


def test_reliable_queue_then_flush():
    outbox._store([])  # clear
    missing_dir = os.path.join(_work(), "cloud")
    path = os.path.join(missing_dir, "log.xlsx")
    status = reliable.excel_append(path, {"Project Number": "9", "Client": "Z"}, label="Log")
    assert status == "queued"
    assert outbox.pending_count() == 1
    # Still unreachable → stays pending.
    assert outbox.flush()["pending"] == 1
    # Bring the location back with a matching workbook, then flush succeeds.
    os.makedirs(missing_dir, exist_ok=True)
    excel.ensure_workbook(path, ["Project Number", "Client"])
    res = outbox.flush()
    assert res["done"] == 1 and res["pending"] == 0
    assert any(str(r["Project Number"]) == "9" for r in excel.read_rows(path))


def test_search_all_columns():
    work = _work()
    path = os.path.join(work, "projects.xlsx")
    excel.ensure_workbook(path, ["Project Number", "Client", "Address", "Contractor"])
    excel.append_row(path, {"Project Number": "100", "Client": "Acme", "Address": "12 Oak St", "Contractor": "BuildCo"})
    cfg = config.default_config()
    cfg["paths"]["project_list"] = path
    assert search.search(cfg, "oak")[0].title.startswith("100")
    assert search.search(cfg, "buildco")            # contractor column
    assert search.search(cfg, "acme")               # client column


def test_search_files_finds_folders_and_files():
    work = _work()
    job = os.path.join(work, "2025-200 - Callahan Residence")
    os.makedirs(os.path.join(job, "05 - Redlines"))
    open(os.path.join(job, "05 - Redlines", "Callahan redline.pdf"), "w").close()
    cfg = config.default_config()
    cfg["base_folder"] = work
    hits = search.search_files(cfg, "callahan")
    names = {h.name for h in hits}
    assert any("Callahan" in n for n in names)
    assert any(h.is_dir for h in hits) and any(not h.is_dir for h in hits)


def test_checklist_scan():
    work = _work()
    folder = os.path.join(work, "job")
    os.makedirs(folder)
    open(os.path.join(folder, "Signed Proposal 2025.pdf"), "w").close()
    needs = checklist.scan_job("1", folder,
                               [{"label": "Signed Proposal", "keywords": ["proposal"]},
                                {"label": "Calcs", "keywords": ["calc"]}],
                               ["Billing rate confirmed"])
    present = {d.label: d.present for d in needs.documents}
    assert present["Signed Proposal"] and not present["Calcs"]
    assert needs.missing_info == ["Billing rate confirmed"]


def test_billing_summary():
    billing.set_contract("B1", 10000, "Test")
    billing.add_entry("B1", {"date": "1/1", "amount": 2500, "hours": 10})
    s = billing.summarize("B1")
    assert s.contract_amount == 10000 and s.billed == 2500 and s.remaining == 7500


def test_jobdata_merge_keeps_values():
    jobdata.update("J1", {"client": "Acme", "address": "12 Oak"})
    jobdata.update("J1", {"client": "", "contractor": "BuildCo"})  # blank must not wipe
    d = jobdata.get("J1")
    assert d["client"] == "Acme" and d["contractor"] == "BuildCo"


def test_activity_counts():
    activity.log(activity.JOB_CREATED, project="1")
    activity.log(activity.EMAIL_SENT, project="1")
    today = activity.counts_today()
    assert today.get(activity.JOB_CREATED, 0) >= 1


def test_fieldscan_value_for():
    jv = {"client": "Acme", "number": "100"}
    assert fieldscan.value_for("client_name", jv) == "Acme"
    assert fieldscan.value_for("project_number", jv) == "100"
    assert fieldscan.value_for("nonexistent", jv) == ""


def test_extract_fields():
    text = "Project No: 2025-104\nClient: Acme Corp\nFee: $12,500.00\nemail joe@acme.com"
    f = extract.extract_fields(text)
    assert f.get("number") == "2025-104"
    assert "acme" in f.get("client", "").lower()
    assert f.get("client_email") == "joe@acme.com"


def test_template_placeholders_and_fill():
    from docx import Document
    work = _work()
    tpl = os.path.join(work, "letter.docx")
    doc = Document()
    doc.add_paragraph("Dear {{ client }}, re project {{ project_number }}.")
    doc.save(tpl)
    keys = templates.get_placeholders(tpl)
    assert "client" in keys and "project_number" in keys
    out = os.path.join(work, "out.docx")
    templates.fill_template(tpl, {"client": "Acme", "project_number": "100"}, out)
    filled = "\n".join(p.text for p in Document(out).paragraphs)
    assert "Acme" in filled and "100" in filled


if __name__ == "__main__":
    fns = [v for k, v in sorted(globals().items()) if k.startswith("test_") and callable(v)]
    passed = 0
    for fn in fns:
        fn()
        print(f"PASS {fn.__name__}")
        passed += 1
    print(f"\n{passed}/{len(fns)} tests passed")
