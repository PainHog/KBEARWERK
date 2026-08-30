"""Job "needed items" tracker.

For each job there is a mandatory list of things that must be on hand - some are
documents that should live in the job folder (proposal, contract, calcs...) and
some are pieces of information someone has to confirm (billing rate, contractor).

This module:
  * scans a job's folder and reports which required **documents** are present or
    missing (by looking for keywords in file names),
  * remembers the manual **info** check-offs per project (local JSON), and
  * builds the combined "what's still missing" list she can print for a meeting
    or hand to an employee.
"""

from __future__ import annotations

import json
import os
from dataclasses import dataclass, field
from typing import Any, Dict, List

from ..config import config_dir


# ---------------------------------------------------------------------------
# Manual check-off storage (per project)
# ---------------------------------------------------------------------------

def _store_path() -> str:
    return os.path.join(config_dir(), "needed.json")


def _load_store() -> Dict[str, Any]:
    path = _store_path()
    if not os.path.exists(path):
        return {}
    try:
        with open(path, "r", encoding="utf-8") as fh:
            data = json.load(fh)
        return data if isinstance(data, dict) else {}
    except (json.JSONDecodeError, OSError):
        return {}


def _save_store(data: Dict[str, Any]) -> None:
    path = _store_path()
    tmp = path + ".tmp"
    with open(tmp, "w", encoding="utf-8") as fh:
        json.dump(data, fh, indent=2)
    os.replace(tmp, path)


def get_info_state(project_number: str) -> Dict[str, bool]:
    """Return {info_label: confirmed?} for a project."""
    return _load_store().get((project_number or "").strip(), {})


def set_info_state(project_number: str, label: str, confirmed: bool) -> None:
    data = _load_store()
    key = (project_number or "").strip()
    rec = data.get(key, {})
    rec[label] = bool(confirmed)
    data[key] = rec
    _save_store(data)


# ---------------------------------------------------------------------------
# Folder scan
# ---------------------------------------------------------------------------

@dataclass
class DocResult:
    label: str
    present: bool
    found: str = ""       # file name that satisfied it


@dataclass
class JobNeeds:
    project_number: str
    folder: str
    documents: List[DocResult] = field(default_factory=list)
    info: List[Dict[str, Any]] = field(default_factory=list)  # {label, confirmed}

    @property
    def missing_documents(self) -> List[DocResult]:
        return [d for d in self.documents if not d.present]

    @property
    def missing_info(self) -> List[str]:
        return [i["label"] for i in self.info if not i["confirmed"]]

    @property
    def all_complete(self) -> bool:
        return not self.missing_documents and not self.missing_info


def _all_filenames(folder: str) -> List[str]:
    names: List[str] = []
    for root, _dirs, files in os.walk(folder):
        for f in files:
            names.append(f.lower())
    return names


def scan_job(
    project_number: str,
    folder: str,
    required_documents: List[Dict[str, Any]],
    required_info: List[str],
) -> JobNeeds:
    """Scan ``folder`` for required documents and combine with manual info state."""
    filenames = _all_filenames(folder) if folder and os.path.isdir(folder) else []

    docs: List[DocResult] = []
    for item in required_documents:
        label = item.get("label", "")
        keywords = [k.lower() for k in item.get("keywords", []) if k]
        found = ""
        for name in filenames:
            if any(k in name for k in keywords):
                found = name
                break
        docs.append(DocResult(label=label, present=bool(found), found=found))

    state = get_info_state(project_number)
    info = [{"label": lbl, "confirmed": bool(state.get(lbl, False))} for lbl in required_info]

    return JobNeeds(project_number=project_number, folder=folder, documents=docs, info=info)


# ---------------------------------------------------------------------------
# Printable "missing items" sheet
# ---------------------------------------------------------------------------

def generate_missing_sheet(dest_path: str, needs: JobNeeds, project_name: str = "", company_name: str = "") -> str:
    """Write a printable Word sheet of everything still missing for a job."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception as exc:  # pragma: no cover
        raise RuntimeError(f"Word engine (python-docx) not available: {exc}")
    import datetime

    os.makedirs(os.path.dirname(os.path.abspath(dest_path)) or ".", exist_ok=True)
    doc = Document()

    if company_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(company_name)
        r.bold = True
        r.font.size = Pt(15)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Outstanding Items")
    tr.bold = True
    tr.font.size = Pt(16)
    tr.font.color.rgb = RGBColor(0xB0, 0x2A, 0x2A)

    hdr = doc.add_paragraph()
    label = f"Job: {needs.project_number}"
    if project_name:
        label += f" — {project_name}"
    hr = hdr.add_run(label)
    hr.bold = True
    hr.font.size = Pt(12)
    doc.add_paragraph(f"Printed: {datetime.date.today().strftime('%m/%d/%Y')}")

    doc.add_paragraph()
    dh = doc.add_paragraph().add_run("Missing Documents")
    dh.bold = True
    dh.font.size = Pt(12)
    if needs.missing_documents:
        for d in needs.missing_documents:
            doc.add_paragraph(d.label, style="List Bullet")
    else:
        doc.add_paragraph("None — all documents present. ✓")

    doc.add_paragraph()
    ih = doc.add_paragraph().add_run("Missing Information")
    ih.bold = True
    ih.font.size = Pt(12)
    if needs.missing_info:
        for lbl in needs.missing_info:
            doc.add_paragraph(lbl, style="List Bullet")
    else:
        doc.add_paragraph("None — all information confirmed. ✓")

    doc.save(dest_path)
    return dest_path
