"""Held-invoice tracking from the Excel invoice ledger.

The office keeps every invoice in an Excel ledger; some are sent, some are HELD.
Held ones are easy to lose track of. This module finds every held invoice, and -
because the engineer's nickname sits in the notes next to "HELD" - can draft a
reminder email to that engineer (resolved to their real address via the contacts
directory).
"""

from __future__ import annotations

import os
from dataclasses import dataclass, field
from typing import Any, Dict, List

from . import excel, people


# Header names we try to read from the ledger (matched loosely).
_INVOICE_KEYS = ["invoice", "invoice #", "invoice no", "invoice number", "inv #", "inv no"]
_CLIENT_KEYS = ["client", "client name", "project", "project name", "job", "owner"]
_AMOUNT_KEYS = ["amount", "total", "balance", "invoice amount", "amount due"]
_STATUS_KEYS = ["status", "state"]
_NOTES_KEYS = ["notes", "note", "comments", "comment", "remarks"]
_ENGINEER_KEYS = ["engineer", "pm", "project manager", "assigned"]
_NUMBER_KEYS = ["project number", "project no", "job number", "job #", "number"]


def _get(row: Dict[str, Any], keys: List[str]) -> str:
    lowered = {str(k).strip().lower(): v for k, v in row.items()}
    for k in keys:
        v = lowered.get(k)
        if v not in (None, ""):
            return str(v).strip()
    return ""


@dataclass
class HeldInvoice:
    invoice: str
    client: str
    amount: str
    engineer: str          # nickname (best guess)
    notes: str
    project_number: str = ""
    row: Dict[str, Any] = field(default_factory=dict)


def _guess_engineer(config: Dict[str, Any], notes: str, explicit: str) -> str:
    """Find the engineer nickname from the notes (next to HELD) or the engineer column."""
    if explicit:
        return explicit
    nicks = people.nicknames(config)
    low = notes.lower()
    for nk in nicks:
        if nk and nk.lower() in low:
            return nk
    # Fall back: the last word of the notes is often the nickname.
    tokens = [t for t in notes.replace("-", " ").split() if t.isalpha()]
    return tokens[-1] if tokens else ""


def load_held(config: Dict[str, Any]) -> List[HeldInvoice]:
    """Return every held invoice from the ledger. Empty if the ledger isn't set."""
    inv_cfg = config.get("invoices", {})
    path = config.get("paths", {}).get("invoice_ledger", "")
    if not path or not os.path.exists(path):
        return []
    sheet = inv_cfg.get("sheet") or None
    try:
        rows = excel.read_rows(path, sheet_name=sheet)
    except Exception:
        return []

    terms = [t.lower() for t in inv_cfg.get("held_terms", ["held"]) if t]
    held: List[HeldInvoice] = []
    for row in rows:
        status = _get(row, _STATUS_KEYS)
        notes = _get(row, _NOTES_KEYS)
        haystack = f"{status} {notes}".lower()
        if not any(t in haystack for t in terms):
            continue
        held.append(HeldInvoice(
            invoice=_get(row, _INVOICE_KEYS),
            client=_get(row, _CLIENT_KEYS),
            amount=_get(row, _AMOUNT_KEYS),
            engineer=_guess_engineer(config, notes, _get(row, _ENGINEER_KEYS)),
            notes=notes,
            project_number=_get(row, _NUMBER_KEYS),
            row=row,
        ))
    return held


def group_by_engineer(held: List[HeldInvoice]) -> Dict[str, List[HeldInvoice]]:
    groups: Dict[str, List[HeldInvoice]] = {}
    for h in held:
        groups.setdefault(h.engineer or "(unknown)", []).append(h)
    return groups


def generate_report(dest_path: str, held: List[HeldInvoice], company_name: str = "") -> str:
    """Write a printable list of every held invoice. Returns the path."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception as exc:  # pragma: no cover
        raise RuntimeError(f"Word engine (python-docx) not available: {exc}")
    import datetime

    os.makedirs(os.path.dirname(os.path.abspath(dest_path)) or ".", exist_ok=True)
    doc = Document()
    if company_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(company_name)
        r.bold = True
        r.font.size = Pt(15)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Held Invoices")
    tr.bold = True
    tr.font.size = Pt(16)
    tr.font.color.rgb = RGBColor(0xB2, 0x6A, 0x00)
    doc.add_paragraph(f"Printed: {datetime.date.today().strftime('%m/%d/%Y')}    "
                      f"Total held: {len(held)}")
    doc.add_paragraph()

    if not held:
        doc.add_paragraph("No held invoices. ✓")
        doc.save(dest_path)
        return dest_path

    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    for i, head in enumerate(["Invoice", "Client", "Amount", "Engineer", "Notes"]):
        run = table.rows[0].cells[i].paragraphs[0].add_run(head)
        run.bold = True
    for h in held:
        cells = table.add_row().cells
        cells[0].text = h.invoice
        cells[1].text = h.client
        cells[2].text = h.amount
        cells[3].text = h.engineer
        cells[4].text = h.notes
    doc.save(dest_path)
    return dest_path
