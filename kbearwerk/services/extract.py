"""Pull information out of a dropped document (proposal, contract, ...).

Reads the text of a PDF or Word file and makes best-effort guesses at the fields
we care about (project number, client, address, contractor, fee, dates,
contacts). She reviews and corrects - extraction is a head-start, never the final
word, and it never blocks her.

The heuristics here are deliberately generic; once we have samples of her real
proposals we tune the patterns to her firm's format for near-perfect pulls.
"""

from __future__ import annotations

import os
import re
from typing import Dict

try:
    from pypdf import PdfReader
    _PDF_OK = True
except BaseException:  # noqa: BLE001 - a broken native dep can raise beyond Exception
    _PDF_OK = False

try:
    from docx import Document
    _DOCX_OK = True
except Exception:  # pragma: no cover
    _DOCX_OK = False


class ExtractError(Exception):
    """A user-friendly extraction problem."""


def extract_text(path: str) -> str:
    """Return the plain text of a PDF, DOCX or TXT file."""
    if not os.path.exists(path):
        raise ExtractError(f"File not found:\n{path}")
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        if not _PDF_OK:
            raise ExtractError("The PDF reader (pypdf) isn't installed.")
        try:
            reader = PdfReader(path)
            return "\n".join((page.extract_text() or "") for page in reader.pages)
        except Exception as exc:
            raise ExtractError(f"Couldn't read the PDF:\n{exc}")
    if ext in (".docx", ".dotx"):
        if not _DOCX_OK:
            raise ExtractError("The Word reader (python-docx) isn't installed.")
        try:
            doc = Document(path)
            parts = [p.text for p in doc.paragraphs]
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
            return "\n".join(parts)
        except Exception as exc:
            raise ExtractError(f"Couldn't read the Word file:\n{exc}")
    if ext in (".txt", ".csv"):
        with open(path, "r", encoding="utf-8", errors="replace") as fh:
            return fh.read()
    raise ExtractError(
        f"I can read PDF, Word (.docx) and text files. This is a {ext or 'unknown'} file."
    )


# --- field heuristics -------------------------------------------------------

_LABEL_PATTERNS = {
    "number": [r"project\s*(?:no\.?|number|#)\s*[:#]?\s*([A-Za-z0-9\-\.]+)",
               r"job\s*(?:no\.?|number|#)\s*[:#]?\s*([A-Za-z0-9\-\.]+)"],
    "name": [r"project\s*name\s*[:\-]?\s*(.+)", r"re\s*[:\-]\s*(.+)"],
    "client": [r"client\s*[:\-]?\s*(.+)", r"owner\s*[:\-]?\s*(.+)",
               r"prepared\s+for\s*[:\-]?\s*(.+)"],
    "contractor": [r"contractor\s*[:\-]?\s*(.+)", r"general\s+contractor\s*[:\-]?\s*(.+)"],
    "address": [r"(?:project\s+)?address\s*[:\-]?\s*(.+)",
                r"site\s+address\s*[:\-]?\s*(.+)",
                r"location\s*[:\-]?\s*(.+)"],
    "fee": [r"(?:fee|contract\s+amount|total\s+fee|lump\s+sum)\s*[:\-]?\s*\$?\s*([\d,]+(?:\.\d{2})?)"],
    "proposal_date": [r"date\s*[:\-]?\s*([A-Za-z0-9,/\.\- ]{6,20})"],
}

_EMAIL_RE = re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}")
_PHONE_RE = re.compile(r"(?:\+?1[\s.\-]?)?\(?\d{3}\)?[\s.\-]?\d{3}[\s.\-]?\d{4}")
_MONEY_RE = re.compile(r"\$\s?([\d,]{3,}(?:\.\d{2})?)")


def _first_match(text: str, patterns) -> str:
    for pat in patterns:
        m = re.search(pat, text, flags=re.IGNORECASE)
        if m:
            val = m.group(1).strip()
            # Trim trailing junk (next label starting, long tails).
            val = re.split(r"\s{2,}|\n", val)[0].strip(" .:-")
            if val:
                return val
    return ""


def extract_fields(text: str) -> Dict[str, str]:
    """Best-effort field extraction. Returns only the fields it found."""
    fields: Dict[str, str] = {}
    for key, patterns in _LABEL_PATTERNS.items():
        val = _first_match(text, patterns)
        if val:
            fields[key] = val

    email = _EMAIL_RE.search(text)
    if email:
        fields["client_email"] = email.group(0)
    phone = _PHONE_RE.search(text)
    if phone:
        fields["client_phone"] = phone.group(0)
    if "fee" not in fields:
        money = _MONEY_RE.search(text)
        if money:
            fields["fee"] = money.group(1)

    return fields


def extract_from_file(path: str) -> Dict[str, str]:
    """Convenience: read a file and return best-guess fields."""
    return extract_fields(extract_text(path))
