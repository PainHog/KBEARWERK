"""Fill-in document templates.

She has blank "form" documents she copies and fills by hand. This turns that into:
choose the template, type into a form, and the app produces a filled copy (and
can save it to several folders at once).

How a template works: we mark the fill-in spots in the blank Word document with
double-brace tags, e.g. ``{{ client }}`` or ``{{ project_address }}``. The app
scans the template, builds a form from those tags automatically, then renders a
filled copy.

We use ``docxtpl`` when available (handles tables/rich layout). If it isn't
installed we fall back to a plain python-docx text replacement so basic templates
still work.
"""

from __future__ import annotations

import os
import re
from typing import Dict, List

_PLACEHOLDER_RE = re.compile(r"{{\s*([a-zA-Z0-9_]+)\s*}}")

try:
    from docxtpl import DocxTemplate
    _DOCXTPL_OK = True
except Exception:  # pragma: no cover
    _DOCXTPL_OK = False

try:
    from docx import Document
    _DOCX_OK = True
except Exception:  # pragma: no cover
    _DOCX_OK = False


class TemplateError(Exception):
    """A user-friendly template problem."""


def humanize(key: str) -> str:
    """Turn a placeholder key like ``project_address`` into ``Project Address``."""
    return re.sub(r"[_\s]+", " ", key).strip().title()


def _scan_docx_text(path: str) -> str:
    if not _DOCX_OK:
        return ""
    try:
        doc = Document(path)
    except Exception as exc:
        raise TemplateError(f"Couldn't open the template:\n{os.path.basename(path)}\n\n({exc})")
    parts: List[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def get_placeholders(path: str) -> List[str]:
    """Return the unique fill-in field keys found in the template, in order."""
    if not os.path.exists(path):
        raise TemplateError(f"Template file not found:\n{path}")
    keys: List[str] = []

    if _DOCXTPL_OK:
        try:
            doc = DocxTemplate(path)
            found = doc.get_undeclared_template_variables()
            keys = sorted(found)
        except Exception:
            keys = []

    if not keys:  # fallback / non-docxtpl
        text = _scan_docx_text(path)
        seen = set()
        for m in _PLACEHOLDER_RE.finditer(text):
            k = m.group(1)
            if k not in seen:
                seen.add(k)
                keys.append(k)
    return keys


def fill_template(template_path: str, context: Dict[str, str], out_path: str) -> str:
    """Render ``template_path`` with ``context`` and save to ``out_path``."""
    if not os.path.exists(template_path):
        raise TemplateError(f"Template file not found:\n{template_path}")
    os.makedirs(os.path.dirname(os.path.abspath(out_path)) or ".", exist_ok=True)

    if _DOCXTPL_OK:
        try:
            doc = DocxTemplate(template_path)
            doc.render(context)
            doc.save(out_path)
            return out_path
        except Exception as exc:
            raise TemplateError(f"Couldn't fill the template:\n{exc}")

    # Fallback: literal text replacement of {{ key }} tokens.
    if not _DOCX_OK:
        raise TemplateError("No Word engine is installed (need docxtpl or python-docx).")
    try:
        doc = Document(template_path)
        for para in doc.paragraphs:
            _replace_in_paragraph(para, context)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _replace_in_paragraph(para, context)
        doc.save(out_path)
        return out_path
    except Exception as exc:
        raise TemplateError(f"Couldn't fill the template:\n{exc}")


def _replace_in_paragraph(para, context: Dict[str, str]) -> None:
    text = para.text
    if "{{" not in text:
        return
    def repl(m):
        return str(context.get(m.group(1), m.group(0)))
    new_text = _PLACEHOLDER_RE.sub(repl, text)
    if new_text != text:
        # Rewrite the whole paragraph in a single run (loses inline formatting,
        # acceptable for the fallback path only).
        for run in list(para.runs):
            run.text = ""
        if para.runs:
            para.runs[0].text = new_text
        else:
            para.add_run(new_text)


def save_copies(src_path: str, dest_dirs: List[str], filename: str) -> List[str]:
    """Copy an already-generated document into each destination folder.

    Returns the list of paths written. Folders that don't exist are created.
    """
    import shutil
    from .files import clean_name, unique_destination

    stem, ext = os.path.splitext(filename)
    filename = clean_name(stem) + ext
    written: List[str] = []
    for d in dest_dirs:
        if not d:
            continue
        dest = unique_destination(d, filename)
        shutil.copy2(src_path, dest)
        written.append(dest)
    return written
