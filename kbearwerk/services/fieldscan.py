"""Work out which fields are still missing for a job, per document.

Compares a job's saved data (the pull-list) against the fields each document /
template needs, and builds a printable report grouped by document, with every
unfilled field flagged. This never blocks anything - it's a to-do list of what's
still needed.
"""

from __future__ import annotations

import os
from dataclasses import dataclass, field
from typing import Any, Dict, List

from . import templates


# Placeholder key -> the job-data keys that can satisfy it.
_ALIASES = {
    "client": ["client", "client_name", "owner"],
    "client_name": ["client", "client_name", "owner"],
    "project": ["name", "project", "project_name"],
    "project_name": ["name", "project_name", "project"],
    "project_number": ["number", "project_number", "job_number"],
    "job_number": ["number", "job_number", "project_number"],
    "number": ["number", "project_number", "job_number"],
    "address": ["address", "project_address", "location"],
    "project_address": ["address", "project_address", "location"],
    "contractor": ["contractor", "general_contractor", "gc"],
    "engineer": ["engineer", "pm"],
    "date": ["date", "date_opened", "proposal_date"],
}


def _norm(s: str) -> str:
    return str(s).strip().lower().replace(" ", "_")


def value_for(placeholder: str, job_values: Dict[str, Any]) -> str:
    """Return the job value that satisfies ``placeholder``, or '' if none."""
    lowered = {_norm(k): v for k, v in job_values.items()}
    candidates = _ALIASES.get(_norm(placeholder), [_norm(placeholder)])
    for c in candidates:
        v = lowered.get(c)
        if v not in (None, "") and str(v).strip():
            return str(v)
    return ""


@dataclass
class DocFields:
    document: str
    missing: List[str] = field(default_factory=list)
    present: List[str] = field(default_factory=list)


def scan(config: Dict[str, Any], job_values: Dict[str, Any]) -> List[DocFields]:
    """For every registered template, list which of its fields are present/missing."""
    groups: List[DocFields] = []
    for tpl in config.get("templates", []):
        path = tpl.get("path", "")
        name = tpl.get("name", os.path.basename(path))
        if not path or not os.path.exists(path):
            continue
        try:
            keys = templates.get_placeholders(path)
        except Exception:
            continue
        g = DocFields(document=name)
        for k in keys:
            if value_for(k, job_values):
                g.present.append(templates.humanize(k))
            else:
                g.missing.append(templates.humanize(k))
        groups.append(g)
    return groups


def generate_report(dest_path: str, project_number: str, project_name: str,
                    groups: List[DocFields], company_name: str = "") -> str:
    """Write a printable 'missing fields' report. Missing fields are red MISSING."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception as exc:  # pragma: no cover
        raise RuntimeError(f"Word engine (python-docx) not available: {exc}")
    import datetime

    os.makedirs(os.path.dirname(os.path.abspath(dest_path)) or ".", exist_ok=True)
    doc = Document()

    if company_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(company_name)
        r.bold = True
        r.font.size = Pt(15)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Missing Fields Report")
    tr.bold = True
    tr.font.size = Pt(16)

    head = doc.add_paragraph()
    label = f"Job: {project_number}"
    if project_name:
        label += f" — {project_name}"
    hr = head.add_run(label)
    hr.bold = True
    hr.font.size = Pt(12)
    doc.add_paragraph(f"Printed: {datetime.date.today().strftime('%m/%d/%Y')}")
    doc.add_paragraph()

    any_missing = False
    for g in groups:
        dh = doc.add_paragraph().add_run(g.document)
        dh.bold = True
        dh.font.size = Pt(13)
        if not g.missing:
            ok = doc.add_paragraph("All fields present").runs[0]
            ok.font.color.rgb = RGBColor(0x1F, 0x7A, 0x44)
        else:
            any_missing = True
            for fieldname in g.missing:
                para = doc.add_paragraph(style="List Bullet")
                para.add_run(f"{fieldname}:  ")
                miss = para.add_run("MISSING")
                miss.bold = True
                miss.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        doc.add_paragraph()

    if not groups:
        doc.add_paragraph("No templates are set up yet, so there are no document "
                          "fields to check. Add templates in the Documents panel.")
    elif not any_missing:
        note = doc.add_paragraph("Nothing missing — every document has all its fields. ✓").runs[0]
        note.font.color.rgb = RGBColor(0x1F, 0x7A, 0x44)

    doc.save(dest_path)
    return dest_path
